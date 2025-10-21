import os
import uvicorn
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import re
from io import BytesIO
import gc
app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

def has_color(run):
    return run.font.color.rgb is not None

def has_highlight_color(run):
    return run.font.highlight_color is not None

def has_bold(run):
    return run.font.bold is not None

def is_correct_answer(run):
    """Kiểm tra run có phải đáp án đúng không (bold, highlight, hoặc màu chữ khác)"""
    return (run.bold is True or 
            run.font.highlight_color is not None or
            (run.font.color and run.font.color.rgb is not None))
    
def clean_question_text(question_text):
    """Xóa các từ thừa trong câu hỏi như 'Câu n:', 'câu n:', 'số:'"""
    # Loại bỏ các pattern thường gặp ở đầu câu hỏi
    patterns_to_remove = [
        r'^Câu\s*\d+[:\.]?\s*',   # Câu 1:, Câu 2., Câu 3
        r'^câu\s*\d+[:\.]?\s*',   # câu 1:, câu 2., câu 3  
        r'^\d+[:\.]?\s*',         # 1:, 2., 3
    ]
    
    cleaned_text = question_text
    for pattern in patterns_to_remove:
        cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE).strip()
    
    return cleaned_text

def clean_answer_text(answer_text):
    """Xóa các ký tự thừa trong đáp án như A., B., a), b), 1., 2."""
    # Loại bỏ các pattern thường gặp ở đầu đáp án
    patterns_to_remove = [
        r'^[a-dA-D][\.\)]\s*',    # A., B., a), b)
        r'^[1-4][\.\)]\s*',       # 1., 2., 1), 2)
    ]
    
    cleaned_text = answer_text
    for pattern in patterns_to_remove:
        cleaned_text = re.sub(pattern, '', cleaned_text).strip()
    
    return cleaned_text

def format_valid_questions(questions):
    """Format và làm sạch các câu hỏi hợp lệ"""
    formatted_questions = []
    
    for question in questions:
        formatted_question = {
             "id": question["id"],
            "question": clean_question_text(question["question"]),
            "answers": [clean_answer_text(answer) for answer in question["answers"]],
            "correct": question["correct"]
        }
        formatted_questions.append(formatted_question)
    
    return formatted_questions

# def process_quiz_format(doc):
#     """Process document with quiz format (1 question, 4 answers)"""
#     data = []
    
#     paragraphs = iter(doc.paragraphs)
#     for paragraph in paragraphs:
#         if paragraph.text.startswith("Câu") or paragraph.text.endswith("?") or paragraph.text.endswith(":"):
#             # Remove numbering pattern like "1. ", "2. ", etc. at the beginning of the question
#             question_text = paragraph.text.strip()
#             question_text = re.sub(r"^\d+\.\s*", "", question_text)
#             # Also remove "Câu XX: " pattern if present
#             question_text = re.sub(r"^Câu\s+\d+\s*[:.\s]\s*", "", question_text)
            
#             question_data = {
#                 "id": len(data) + 1,  # Auto-increment ID
#                 "question": question_text,
#                 "answers": [],
#                 "correct": ""
#             }

#             options = []
#             correct_answer = ""

#             for i in range(4):  # Assuming 4 options for each question
#                 try:
#                     option_paragraph = next(paragraphs)
#                     if not option_paragraph.text.strip():  # Skip empty paragraphs
#                         continue
                        
#                     # Clean up option text (remove A., B., etc. prefixes)
#                     option_text = option_paragraph.text.strip()
#                     option_text = re.sub(r"^[A-D]\.\s*", "", option_text)
                    
#                     options.append(option_text)
                    
#                     # Check if this option is marked as correct (bold or highlighted)
#                     is_correct = False
#                     for option_run in option_paragraph.runs:
#                         if has_bold(option_run) or has_highlight_color(option_run) or has_color(option_run):
#                             correct_answer = str(i)  # Use index position (0, 1, 2, 3)
#                             is_correct = True
#                             break
                            
#                 except StopIteration:
#                     break

#             question_data["answers"] = options
#             question_data["correct"] = correct_answer
#             data.append(question_data)
    
#     return data

def process_quiz_format(doc):
    """Process document with Q&A format (1 question, 1 answer)"""
    quiz_data = []
    current_question = None
    
    # Patterns để nhận diện câu hỏi và đáp án
    question_patterns = [
        r'^(Câu\s*\d+[:\.]?)\s*(.*)',  # Câu 1., Câu 2:
        r'^(\d+[:\.])\s*(.*)',          # 1., 2:
        r'^(.*)\?$'                     # Kết thúc bằng ?
    ]
    
    answer_pattern = r'^([a-dA-D][\.\)])\s*(.*)'  # a., B), c.
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Tách paragraph thành nhiều dòng nếu có xuống dòng
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Kiểm tra xem có phải câu hỏi không
            is_question = any(re.match(pattern, line, re.IGNORECASE) for pattern in question_patterns)
            
            if is_question:
                # Lưu câu hỏi trước đó nếu có
                if current_question:
                    quiz_data.append(current_question)
                
                # Tạo câu hỏi mới
                current_question = {
                    "id": len(quiz_data) + 1,  # Auto-increment ID
                    "question": line,
                    "answers": [],
                    "correct": -1  # -1 nghĩa là chưa có đáp án đúng
                }
            
            # Kiểm tra xem có phải đáp án không
            elif current_question:
                match = re.match(answer_pattern, line, re.DOTALL)
                if match:
                    answer_text = match.group(2).strip()
                    current_question["answers"].append(line)
                    
                    # Kiểm tra format của đáp án để xem có phải đáp án đúng không
                    for run in para.runs:
                        if answer_text in run.text and is_correct_answer(run):
                            # Lưu index của đáp án đúng (bắt đầu từ 0)
                            current_question["correct"] = len(current_question["answers"]) - 1
                            break
    
    # Thêm câu hỏi cuối cùng
    if current_question:
        quiz_data.append(current_question)
    
    # Sắp xếp câu hỏi: câu hỏi có đáp án và correct lên trước, câu hỏi thiếu thông tin xuống cuối
    valid_questions = []
    invalid_questions = []
    
    for question in quiz_data:
        # Kiểm tra câu hỏi có đầy đủ thông tin không (có answers và correct >= 0)
        if question["question"] and question["answers"] and question["correct"] >= 0:
            valid_questions.append(question)
        else:
            # Chuẩn hóa invalid questions - đảm bảo có 4 đáp án trống
            invalid_question = {
                "id": question["id"],
                "question": question["question"],
                "answers": question["answers"] if question["answers"] else ["", "", "", ""],
                "correct": question["correct"]
            }
            # Nếu có ít hơn 4 đáp án, thêm đáp án trống
            while len(invalid_question["answers"]) < 4:
                invalid_question["answers"].append("")
            invalid_questions.append(invalid_question)
    
    # Format và làm sạch các câu hỏi hợp lệ
    formatted_valid_questions = format_valid_questions(valid_questions)
    
    # Kết hợp lại: câu hỏi hợp lệ đã format trước, câu hỏi thiếu thông tin sau
    sorted_quiz_data = formatted_valid_questions + invalid_questions
    
    return sorted_quiz_data

@app.get("/")
async def read_root():
    return {
        "message": "DOCX to JSON Converter API",
        "endpoints": {
            "/quiz": "POST - Convert DOCX to quiz format (1 question, 4 answers)",
            "/api/qa": "POST - Convert DOCX to Q&A format (1 question, 1 answer)"
        },
        "usage": "Upload .docx file using multipart/form-data"
    }

@app.post('/quiz')
async def convert_quiz(file: UploadFile = File(...)):
    """API endpoint for quiz format (1 question, 4 answers with correct answer marked)"""
    try:
        # Validate file type
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="File must be a .docx file")
        
        # Read file content into memory
        file_content = await file.read()
        
        # Process document from memory
        doc = Document(BytesIO(file_content))
        result = process_quiz_format(doc)
        
        # Clear memory
        del file_content
        del doc
        gc.collect()
        
        return JSONResponse(content=result)
        
    except Exception as e:
        # Clear memory in case of error
        gc.collect()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

# @app.post('/api/qa')
# async def convert_qa(file: UploadFile = File(...)):
#     """API endpoint for Q&A format (1 question, 1 answer)"""
#     try:
#         # Validate file type
#         if not file.filename.endswith('.docx'):
#             raise HTTPException(status_code=400, detail="File must be a .docx file")
        
#         # Read file content into memory
#         file_content = await file.read()
        
#         # Process document from memory
#         doc = Document(BytesIO(file_content))
#         result = process_qa_format(doc)
        
#         # Clear memory
#         del file_content
#         del doc
#         gc.collect()
        
#         return JSONResponse(content=result)
        
#     except Exception as e:
#         # Clear memory in case of error
#         gc.collect()
#         raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080)) # Mặc định là 8000 cho FastAPI
    uvicorn.run(app, host="0.0.0.0", port=port)